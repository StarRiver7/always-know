from typing import List, Dict, Any
import pdfplumber
from docx import Document
from io import BytesIO
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_community.vectorstores import Milvus
from app.core.config import settings
import logging

logger = logging.getLogger(__name__)


class DocumentService:
    """基于LangChain的文档处理服务类"""

    def __init__(self):
        """初始化LangChain组件"""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=[
                "\n\n",
                "\n",
                "。", "！", "？",
                ".", "!", "?",
                "；", ";",
                "，", ",",
                " ",
                ""
            ]
        )

        self.embeddings = OpenAIEmbeddings(
            api_key=settings.OPENAI_API_KEY,
            base_url=settings.OPENAI_API_BASE,
            model=settings.OPENAI_EMBEDDING_MODEL
        )

        self.chat_model = ChatOpenAI(
            api_key=settings.OPENAI_API_KEY,
            base_url=settings.OPENAI_API_BASE,
            model=settings.OPENAI_CHAT_MODEL,
            temperature=0.7,
            max_tokens=1000
        )

        self.vector_store = None

        logger.info("DocumentService初始化完成")

    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """从PDF文件中提取文本"""
        text = ""
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """从DOCX文件中提取文本"""
        doc = Document(BytesIO(file_bytes))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text

    def extract_text_from_txt(self, file_bytes: bytes) -> str:
        """从TXT文件中提取文本"""
        return file_bytes.decode("utf-8")

    def split_text_into_chunks(self, text: str) -> List[str]:
        """使用LangChain分块文本"""
        return self.text_splitter.split_text(text)

    async def init_vector_store(self):
        """初始化向量存储"""
        if self.vector_store is None:
            self.vector_store = Milvus(
                embedding_function=self.embeddings,
                collection_name=settings.MILVUS_COLLECTION_NAME,
                connection_args={
                    "host": settings.MILVUS_HOST,
                    "port": settings.MILVUS_PORT
                },
                drop_old=True,
                vector_field="embedding",
                text_field="page_content",
                auto_id=True
            )
        return self.vector_store

    async def process_document(self, document_id: int, file_bytes: bytes, filename: str, metadata: Dict = None) -> int:
        """处理文档并存储到知识库"""
        logger.info(f"开始处理文档: document_id={document_id}, filename={filename}, 文件大小={len(file_bytes)} 字节")

        try:
            if filename.endswith(".pdf"):
                logger.info("步骤1: 从PDF提取文本...")
                text = self.extract_text_from_pdf(file_bytes)
            elif filename.endswith(".docx"):
                logger.info("步骤1: 从DOCX提取文本...")
                text = self.extract_text_from_docx(file_bytes)
            elif filename.endswith(".txt"):
                logger.info("步骤1: 从TXT提取文本...")
                text = self.extract_text_from_txt(file_bytes)
            else:
                raise ValueError(f"不支持的文件类型: {filename}")

            logger.info(f"步骤1完成: 文本长度={len(text)} 字符")

            logger.info("步骤2: 使用LangChain分割文本...")
            chunks = self.split_text_into_chunks(text)
            logger.info(f"步骤2完成: 创建了 {len(chunks)} 个块")

            logger.info("步骤3: 创建LangChain Document对象...")
            langchain_docs = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    "document_id": document_id,
                    "chunk_id": i,
                    **(metadata or {})
                }
                langchain_docs.append(LangChainDocument(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            logger.info(f"步骤3完成: 创建了 {len(langchain_docs)} 个Document对象")

            logger.info("步骤4: 插入Milvus向量存储...")
            await self.init_vector_store()
            await self.vector_store.aadd_documents(langchain_docs)
            logger.info("步骤4完成: 向量插入成功")

            logger.info(f"文档处理完成: {len(chunks)} 个块已处理")
            return len(chunks)
        except Exception as e:
            logger.error(f"文档处理失败: {str(e)}", exc_info=True)
            raise


    async def query_knowledge(self, query: str, top_k: int = 5, document_ids: List[int] = None) -> Dict[str, Any]:
        """查询知识库"""
        logger.info(f"开始知识库查询: query={query[:50]}..., top_k={top_k}, document_ids={document_ids}")

        await self.init_vector_store()

        search_kwargs = {"k": top_k}
        if document_ids:
            doc_id_strs = [str(doc_id) for doc_id in document_ids]
            filter_expr = f"document_id in [{','.join(doc_id_strs)}]"
            search_kwargs["expr"] = filter_expr
            logger.info(f"使用filter表达式: {filter_expr}")

        retriever = self.vector_store.as_retriever(
            search_kwargs=search_kwargs
        )

        docs = await retriever.ainvoke(query)
        logger.info(f"检索到的文档数量: {len(docs)}")
        for i, doc in enumerate(docs):
            logger.info(f"文档 {i}: page_content长度={len(doc.page_content) if hasattr(doc, 'page_content') else 'N/A'}")
            logger.info(f"文档 {i} 元数据: {doc.metadata}")
            logger.info(f"文档 {i} 内容预览: {doc.page_content[:100]}...")

        context = "\n\n".join([doc.page_content for doc in docs])

        prompt = PromptTemplate(
            template="""你是一个专业的知识库助手。请根据以下上下文信息回答用户的问题。
如果上下文不足以回答问题，请诚实地说明，不要编造信息。

上下文信息：
{context}

问题：{question}

回答：""",
            input_variables=["context", "question"]
        )

        formatted_prompt = prompt.format(context=context, question=query)
        logger.info(f"发送给LLM的prompt长度: {len(formatted_prompt)}")

        chat_result = await self.chat_model.ainvoke(formatted_prompt)
        answer = chat_result.content if hasattr(chat_result, 'content') else str(chat_result)

        logger.info(f"知识库查询完成，回答长度={len(answer)}")

        sources = []
        for i, doc in enumerate(docs):
            metadata = doc.metadata
            sources.append({
                "document_id": metadata.get("document_id"),
                "chunk_id": metadata.get("chunk_id"),
                "content": doc.page_content,
                "metadata": {k: v for k, v in metadata.items() if k not in ["document_id", "chunk_id"]},
                "score": 1.0
            })

        return {
            "answer": answer,
            "sources": sources
        }

    async def delete_document(self, document_id: int):
        """删除文档"""
        logger.info(f"正在删除文档 {document_id}...")

        # 初始化向量存储
        await self.init_vector_store()
        
        # 构建删除表达式
        expr = f"document_id == {document_id}"
        logger.info(f"使用表达式删除向量: {expr}")
        
        # 执行删除操作
        await self.vector_store.adelete(expr=expr)
        logger.info(f"删除文档 {document_id} 的所有向量成功")

document_service = DocumentService()